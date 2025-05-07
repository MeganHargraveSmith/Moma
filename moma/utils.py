import os
import time
from django.core.exceptions import ValidationError
from django.conf import settings
from PIL import Image
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import img2pdf
import io
import mimetypes



def sanitize_filename(filename):
    """
    Sanitize the filename by removing spaces and adding a timestamp.
    """
    filename = os.path.basename(filename)
    filename = filename.replace(' ', '_')
    timestamp = int(time.time())
    name, ext = os.path.splitext(filename)
    return f"{name}_{timestamp}{ext}"

def validate_file_size(value):
    """
    Validate that the uploaded file size does not exceed the maximum allowed size.
    Raise a ValidationError if the file is too large.
    """
    if value.size > settings.MAX_UPLOAD_SIZE:
        raise ValidationError(f'File size exceeds maximum allowed size of {settings.MAX_UPLOAD_SIZE // 1024 // 1024}MB')

def delete_file_if_exists(file_path):
    """
    Delete the specified file if it exists.
    """
    if file_path and os.path.isfile(file_path):
        os.remove(file_path)

def convert_to_pdf(file_path, original_filename):
    """
    Convert the specified file to PDF format.
    Raise a ValidationError if the file type is unsupported or processing fails.
    """
    if not original_filename:
        raise ValidationError('Original filename is required')
        
    pdf_filename = os.path.splitext(original_filename)[0] + '.pdf'
    pdf_path = os.path.join(settings.MEDIA_ROOT, 'converted_pdfs', pdf_filename)
    
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    
    _, ext = os.path.splitext(original_filename.lower())
    
    try:
        if not ext:
            content_type, _ = mimetypes.guess_type(file_path)
            if content_type == 'application/pdf':
                ext = '.pdf'
            elif content_type in ['image/jpeg', 'image/jpg']:
                ext = '.jpg'
            elif content_type == 'image/png':
                ext = '.png'
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                ext = '.docx'
            else:
                raise ValidationError('Could not determine file type')
                
        if ext == '.pdf':
            with open(file_path, 'rb') as src, open(pdf_path, 'wb') as dst:
                dst.write(src.read())
        elif ext in ['.jpg', '.jpeg', '.png']:
            with open(pdf_path, 'wb') as f:
                f.write(img2pdf.convert(file_path))
        elif ext == '.docx':
            doc = Document(file_path)
            c = canvas.Canvas(pdf_path, pagesize=letter)
            width, height = letter
            
            left_margin = 50
            right_margin = 50
            top_margin = 50
            bottom_margin = 50
            
            usable_width = width - left_margin - right_margin
            usable_height = height - top_margin - bottom_margin
            
            y = height - top_margin
            line_height = 14
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    y -= line_height
                    continue
                
                words = para.text.split()
                current_line = []
                current_line_width = 0
                
                for word in words:
                    word_width = c.stringWidth(word + ' ', 'Helvetica', 12)
                    if current_line_width + word_width <= usable_width:
                        current_line.append(word)
                        current_line_width += word_width
                    else:
                        if current_line:
                            c.drawString(left_margin, y, ' '.join(current_line))
                            y -= line_height
                        
                        current_line = [word]
                        current_line_width = word_width
                
                if current_line:
                    c.drawString(left_margin, y, ' '.join(current_line))
                    y -= line_height
                
                if y < bottom_margin:
                    c.showPage()
                    y = height - top_margin
            
            for table in doc.tables:
                y -= line_height
                
                for row in table.rows:
                    if y < bottom_margin:
                        c.showPage()
                        y = height - top_margin
                    
                    x = left_margin
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            c.drawString(x, y, cell_text)
                        x += 100
                    
                    y -= line_height
                
                y -= line_height
            
            c.save()
        else:
            raise ValidationError(f'Unsupported file type: {ext}')
        
        return pdf_path
        
    except Exception as e:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        raise ValidationError(f'Failed to process file: {str(e)}') 